from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\Asus\OneDrive\Dokumen\Semester 6\MPTI\CashDig-4SR")
OUT = ROOT / "docs" / "portfolio_project"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Laporan_Portofolio_Project_Sikasir_4SR.docx"
ERD = OUT / "erd_sikasir_4sr.png"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(80, 80, 80)


def shade(cell, fill="E8EEF5"):
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def txt(cell, value, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(value))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def borders(table):
    pr = table._tbl.tblPr
    b = pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            b.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "A6A6A6")


def widths(table, col_widths):
    for row in table.rows:
        for i, width in enumerate(col_widths):
            cell = row.cells[i]
            cell.width = Inches(width)
            pr = cell._tc.get_or_add_tcPr()
            tcw = pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                pr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(int(width * 1440)))


def table(doc, headers, rows, col_widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        txt(t.rows[0].cells[i], h, True, size)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            txt(cells[i], value, False, size)
    if col_widths:
        widths(t, col_widths)
    borders(t)
    doc.add_paragraph()
    return t


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def num(doc, text):
    doc.add_paragraph(text, style="List Number")


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def codeblock(doc, title, code, note):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK
    for line in code.strip().splitlines():
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Inches(0.25)
        q.paragraph_format.space_after = Pt(0)
        rr = q.add_run(line)
        rr.font.name = "Consolas"
        rr._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        rr.font.size = Pt(8.3)
    doc.add_paragraph(note)


def draw_erd():
    img = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(img)
    try:
        ft = ImageFont.truetype("arial.ttf", 32)
        fh = ImageFont.truetype("arialbd.ttf", 24)
        f = ImageFont.truetype("arial.ttf", 21)
        fs = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        ft = fh = f = fs = ImageFont.load_default()

    d.text((60, 35), "ERD Sikasir-4SR (Laravel POS Rumah Makan)", fill="#0B2545", font=ft)
    entities = {
        "users": (70, 120, 430, 315, ["id (PK)", "name", "email (unique)", "password", "remember_token"]),
        "categories": (70, 415, 430, 585, ["id (PK)", "name", "description", "timestamps"]),
        "menu_items": (575, 330, 970, 585, ["id (PK)", "category_id (FK)", "name", "price", "stock", "image_url", "is_available", "description"]),
        "tables": (70, 720, 430, 880, ["id (PK)", "number", "qr_code", "timestamps"]),
        "orders": (575, 710, 970, 930, ["id (PK)", "table_id (FK)", "status", "total_price", "note", "timestamps"]),
        "order_items": (1190, 625, 1640, 895, ["id (PK)", "order_id (FK)", "menu_item_id (FK)", "quantity", "subtotal", "timestamps"]),
        "sessions": (1190, 120, 1640, 330, ["id (PK)", "user_id (FK nullable)", "ip_address", "user_agent", "payload", "last_activity"]),
    }
    for name, (x1, y1, x2, y2, fields) in entities.items():
        d.rounded_rectangle([x1, y1, x2, y2], radius=16, outline="#1F4D78", width=3, fill="#F7FAFC")
        d.rectangle([x1, y1, x2, y1 + 42], fill="#E8EEF5", outline="#1F4D78", width=2)
        d.text((x1 + 16, y1 + 9), name, fill="#0B2545", font=fh)
        yy = y1 + 56
        for field in fields:
            d.text((x1 + 18, yy), field, fill="#222222", font=f)
            yy += 28

    def arrow(start, end, label):
        d.line([start, end], fill="#444444", width=3)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            d.polygon([(ex, ey), (ex - 16, ey - 8), (ex - 16, ey + 8)], fill="#444444")
        elif ex < sx:
            d.polygon([(ex, ey), (ex + 16, ey - 8), (ex + 16, ey + 8)], fill="#444444")
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        d.rounded_rectangle([mx - 80, my - 18, mx + 80, my + 18], radius=8, fill="white", outline="#D0D7DE")
        d.text((mx - 32, my - 12), label, fill="#333333", font=fs)

    arrow((430, 500), (575, 455), "1 : N")
    arrow((430, 800), (575, 820), "1 : N")
    arrow((970, 820), (1190, 760), "1 : N")
    arrow((970, 500), (1190, 730), "1 : N")
    arrow((430, 210), (1190, 225), "0..1 : N")
    d.text(
        (60, 1040),
        "Catatan: ERD disusun dari model/controller aktif. Migration perlu diselaraskan karena masih ada kolom lama seperti category, is_active, subtotal, total, dan line_total.",
        fill="#7A5A00",
        font=fs,
    )
    img.save(ERD)


def styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, RGBColor(0, 0, 0), 0, 6),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "Laporan Portofolio Project - Sikasir-4SR"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.text = "CashDig-4SR / MPTI"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def main():
    draw_erd()
    doc = Document()
    styles(doc)

    logo = ROOT / "public" / "assets" / "logo-4sr.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.35))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("LAPORAN PORTOFOLIO PROJECT")
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SISTEM KASIR RUMAH MAKAN 4SR (SIKASIR-4SR)")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Portfolio berdasarkan aktivitas proyek sampai pertemuan 07: Front-end, API/route, dan Back-end Laravel.").italic = True

    table(
        doc,
        ["Komponen", "Keterangan"],
        [
            ["Judul project", "Sikasir-4SR / CashDig-4SR"],
            ["Studi kasus", "Sistem kasir, pemesanan, stok, dan laporan penjualan Rumah Makan 4SR"],
            ["Framework", "Laravel 12, PHP 8.2, Blade, MySQL, JavaScript, CSS"],
            ["Repository", "https://github.com/arico18038/RPL-Project.git"],
            ["User interface", "Kasir, barang & stok, riwayat, laporan, pengaturan, login"],
            ["Tanggal dokumen", date.today().strftime("%d %B %Y")],
        ],
        [1.7, 4.7],
        9.5,
    )

    doc.add_heading("Identitas Tim", level=2)
    table(
        doc,
        ["No", "NIM", "Nama", "Kelas", "JobDesc / Peran"],
        [
            ["1", "2300018038", "ARICO SANGAPTA MELIALA", "[ISI KELAS]", "Konsep sistem, kebutuhan fitur, ERD, relasi database, alur menu-pesanan, validasi kebutuhan, dan koordinasi progres."],
            ["2", "2300018040", "RENO SAPUTRA", "[ISI KELAS]", "Front-end/UI: halaman menu/customer, login admin, dashboard, barang dan stok, riwayat, laporan, dan pengaturan."],
            ["3", "2300018044", "MUSTOFA ABDURRAHIM", "[ISI KELAS]", "Back-end Laravel, database MySQL/MariaDB, route, controller, model, login admin, CRUD menu, stok, pesanan, dan dokumentasi teknis."],
        ],
        [0.45, 1.1, 1.8, 0.9, 2.15],
        8.5,
    )
    doc.add_paragraph("Catatan: tebalkan dan kapitalisasi NIM & NAMA mahasiswa yang bersangkutan sebelum dicetak sesuai ketentuan tugas.")
    doc.add_page_break()

    doc.add_heading("Daftar Isi Ringkas", level=1)
    for item in [
        "1. Halaman Cover dan Identitas Tim",
        "2. Halaman Daftar Pembagian Tugas / JobDesc",
        "3. Ruang Lingkup / Deskripsi Sistem",
        "4. Daftar Seluruh Spesifikasi Kebutuhan (SRS)",
        "5. Isi Laporan Sesuai JobDesc Masing-Masing Anggota",
        "6. Lampiran Pendukung: Pengujian, Temuan, dan Struktur File",
    ]:
        bullet(doc, item)

    doc.add_heading("2. Halaman Daftar Pembagian Tugas / JobDesc", level=1)
    table(
        doc,
        ["No", "Pekan ke-", "List Tugas", "PIC"],
        [
            ["1", "1", "Analisis studi kasus rumah makan, identifikasi aktor admin/kasir/pelanggan, penyusunan ruang lingkup, dan validasi kebutuhan awal.", "Arico Sangapta Meliala"],
            ["2", "2", "Penyusunan konsep sistem, alur menu-pesanan, rancangan awal ERD, serta pembagian kebutuhan fitur ke Front-end dan Back-end.", "Arico Sangapta Meliala"],
            ["3", "3", "Perancangan UI berdasarkan referensi Figma: halaman customer/menu, login admin, dashboard, barang dan stok, riwayat, laporan, dan pengaturan.", "Reno Saputra"],
            ["4", "4", "Implementasi tampilan halaman kasir dan admin: daftar produk, kategori, pencarian, keranjang, ringkasan pembayaran, dan komponen visual dashboard.", "Reno Saputra"],
            ["5", "5", "Pengembangan Back-end Laravel: route, controller, model, autentikasi admin, CRUD menu, validasi stok, dan penyimpanan pesanan.", "Mustofa Abdurrahim"],
            ["6", "6", "Pengelolaan database MySQL/MariaDB, relasi tabel, fitur stok, fitur pesanan, API/route, dan penyesuaian alur transaksi.", "Mustofa Abdurrahim"],
            ["7", "7", "Pengecekan kesesuaian fitur, dokumentasi progres, portofolio, FR/NFR, ERD final, hasil pengujian, dan rekomendasi perbaikan.", "Arico, Reno, Mustofa"],
        ],
        [0.45, 0.75, 4.1, 1.25],
        8.5,
    )

    doc.add_heading("3. Ruang Lingkup / Deskripsi Sistem", level=1)
    doc.add_paragraph(
        "Sikasir-4SR adalah aplikasi kasir berbasis web untuk membantu Rumah Makan 4SR mengelola transaksi penjualan, "
        "ketersediaan menu, stok barang, dan laporan penjualan. Aplikasi dibangun dengan Laravel sebagai backend MVC, "
        "Blade sebagai templating UI, MySQL sebagai basis data, serta JavaScript untuk interaksi keranjang belanja."
    )
    doc.add_heading("Aktor Sistem", level=2)
    for item in [
        "Pengunjung/Pelanggan: melihat menu dan membuat pesanan berdasarkan meja.",
        "Kasir: memilih menu, menambah item ke keranjang, memilih nomor meja, dan mengirim pesanan.",
        "Admin: login, mengelola barang/menu, melihat pesanan, memproses pesanan, melihat riwayat transaksi, dan membuka laporan.",
        "Sistem: melakukan validasi input, menghitung pajak, menyimpan transaksi, dan mengurangi stok secara otomatis.",
    ]:
        bullet(doc, item)
    doc.add_heading("Batasan Sistem", level=2)
    for item in [
        "Metode pembayaran pada UI saat ini difokuskan pada Tunai.",
        "Diskon tersedia pada tampilan kasir, namun perhitungan backend saat ini belum menyimpan nilai diskon sebagai kolom transaksi.",
        "Pengelolaan user masih memakai autentikasi Laravel standar dan belum memiliki halaman manajemen user khusus.",
        "Database lokal memakai MySQL; koneksi perlu aktif agar migrate/status dan proses transaksi berjalan.",
    ]:
        bullet(doc, item)

    doc.add_heading("4. Daftar Seluruh Spesifikasi Kebutuhan (SRS) Sistem", level=1)
    doc.add_heading("4.1 Kebutuhan Fungsional (FR)", level=2)
    table(
        doc,
        ["Kode", "Kebutuhan Fungsional", "Implementasi"],
        [
            ["FR-01", "Sistem menampilkan daftar menu aktif beserta kategori, harga, gambar, dan stok.", "PosController, MenuItem, Category, pos/index.blade.php"],
            ["FR-02", "Pengguna dapat mencari dan memfilter menu berdasarkan kategori.", "public/assets/js/pos.js"],
            ["FR-03", "Kasir dapat menambahkan, mengurangi, dan menghapus item pada keranjang.", "public/assets/js/pos.js"],
            ["FR-04", "Sistem menghitung subtotal, PPN 11%, dan total pembayaran pada UI.", "public/assets/js/pos.js"],
            ["FR-05", "Sistem memvalidasi nomor meja dan item pesanan sebelum menyimpan order.", "OrderController@store"],
            ["FR-06", "Sistem menolak pesanan jika stok habis atau kuantitas melebihi stok tersedia.", "OrderController@store"],
            ["FR-07", "Sistem menyimpan order dan order item dalam transaksi database.", "DB::transaction, Order, OrderItem"],
            ["FR-08", "Sistem mengurangi stok menu setelah order berhasil dibuat.", "MenuItem decrement stock"],
            ["FR-09", "Admin dapat login dan logout dengan sesi Laravel.", "AuthController, middleware auth/guest"],
            ["FR-10", "Admin dapat menambah, mengubah, dan menghapus barang/menu.", "AdminController CRUD menu"],
            ["FR-11", "Admin dapat melihat daftar pesanan dan mengubah status menjadi diproses atau selesai.", "AdminController status order"],
            ["FR-12", "Admin dapat melihat riwayat transaksi dan laporan ringkas penjualan.", "AdminController history/report"],
        ],
        [0.65, 3.75, 2.05],
        8,
    )

    doc.add_heading("4.2 Kebutuhan Non-Fungsional (NFR)", level=2)
    table(
        doc,
        ["Kode", "Kategori", "Deskripsi"],
        [
            ["NFR-01", "Keamanan", "Halaman admin wajib dilindungi middleware auth; password user disimpan dalam bentuk hash Laravel."],
            ["NFR-02", "Integritas data", "Penyimpanan order harus memakai transaksi database dan lockForUpdate agar stok tidak minus pada akses bersamaan."],
            ["NFR-03", "Usability", "UI kasir harus cepat digunakan: tombol tambah, ringkasan pembayaran, shortcut F2/F7/F9, dan status stok terlihat jelas."],
            ["NFR-04", "Maintainability", "Kode dipisah dalam MVC Laravel: route, controller, model, migration, seeder, view, CSS, dan JavaScript."],
            ["NFR-05", "Reliability", "Validasi backend tetap dilakukan meskipun validasi UI dilewati."],
            ["NFR-06", "Performance", "Query utama memakai eager loading relasi category/order_items/table untuk mengurangi N+1 query."],
            ["NFR-07", "Compatibility", "Aplikasi berjalan pada PHP 8.2, Laravel 12, MySQL, dan browser modern."],
            ["NFR-08", "Auditability", "Order menyimpan waktu pembuatan melalui timestamps sehingga riwayat transaksi dapat dilacak."],
        ],
        [0.65, 1.35, 4.45],
        8.5,
    )

    doc.add_heading("4.3 ERD dan Perancangan Basis Data", level=2)
    doc.add_paragraph(
        "ERD berikut disusun berdasarkan model dan controller aktif pada aplikasi. Entitas utama yang mendukung proses POS adalah "
        "users, categories, menu_items, tables, orders, dan order_items. Tabel sessions digunakan oleh Laravel untuk sesi login."
    )
    doc.add_picture(str(ERD), width=Inches(6.45))
    caption(doc, "Gambar 1. Entity Relationship Diagram Sikasir-4SR")
    table(
        doc,
        ["Entitas", "Fungsi", "Relasi Utama"],
        [
            ["users", "Menyimpan akun admin/kasir untuk login.", "Satu user dapat memiliki banyak session."],
            ["categories", "Mengelompokkan menu seperti makanan/minuman.", "Satu category memiliki banyak menu_items."],
            ["menu_items", "Menyimpan data menu/barang: nama, harga, stok, gambar, status aktif.", "Banyak menu_items berada pada satu category; satu menu dapat muncul di banyak order_items."],
            ["tables", "Menyimpan nomor meja dan QR code.", "Satu table dapat memiliki banyak orders."],
            ["orders", "Menyimpan transaksi utama: meja, status, total harga, catatan.", "Satu order memiliki banyak order_items."],
            ["order_items", "Menyimpan detail item yang dipesan, quantity, dan subtotal.", "Banyak order_items mengacu ke satu order dan satu menu_item."],
        ],
        [1.15, 2.75, 2.55],
        8.5,
    )

    doc.add_heading("4.4 Arsitektur Sistem dan Route/API", level=2)
    doc.add_paragraph(
        "Aplikasi memakai pola MVC Laravel. Browser mengakses route di routes/web.php, request diproses controller, "
        "controller membaca/menulis data melalui Eloquent model, kemudian hasil dikirim ke Blade view. JavaScript di "
        "public/assets/js/pos.js menangani interaksi keranjang sebelum form dikirim ke backend."
    )
    table(
        doc,
        ["Method", "URL", "Nama Route", "Fungsi"],
        [
            ["GET", "/", "pos.index", "Halaman kasir utama dan daftar menu."],
            ["GET/POST", "/login", "login/login.store", "Form login dan proses autentikasi."],
            ["POST", "/orders", "orders.store", "Menyimpan pesanan baru."],
            ["GET", "/admin", "admin.index", "Daftar barang dan stok."],
            ["POST/PUT/DELETE", "/admin/menu", "admin.menu.*", "CRUD data menu/barang."],
            ["GET", "/admin/orders", "admin.orders", "Daftar pesanan."],
            ["PATCH", "/admin/orders/{order}/process", "admin.orders.process", "Mengubah status ke processing."],
            ["PATCH", "/admin/orders/{order}/complete", "admin.orders.complete", "Mengubah status ke completed."],
            ["GET", "/admin/riwayat", "admin.history", "Riwayat transaksi."],
            ["GET", "/admin/laporan", "admin.report", "Laporan penjualan."],
        ],
        [0.75, 1.35, 1.6, 2.75],
        8,
    )

    doc.add_heading("5. Isi Laporan Berikut Sesuai JobDesc Masing-Masing", level=1)
    doc.add_paragraph(
        "Bagian ini dibuat sesuai ketentuan: dalam satu tim, tiap anggota tidak ada yang sama. "
        "Setiap anggota memiliki tiga komponen wajib, yaitu: "
        "a. daftar/list pekerjaan sesuai jobdesc; "
        "b. hasil kerja berupa capture-capture dan diberi penjelasan yang sesuai; "
        "c. hasil kerja berupa coding/skrip dan diberi penjelasan yang sesuai."
    )

    doc.add_heading("5.1 ARICO SANGAPTA MELIALA - Analisis Sistem, ERD, dan Validasi Kebutuhan", level=2)
    doc.add_heading("a. Daftar/List Pekerjaan Sesuai JobDesc", level=3)
    for item in [
        "Menyusun konsep sistem Sikasir-4SR sesuai studi kasus rumah makan.",
        "Menentukan kebutuhan fitur utama: menu, pesanan, stok, login admin, riwayat transaksi, dan laporan.",
        "Membantu rancangan ERD, relasi database, serta alur data dari menu ke pesanan dan detail pesanan.",
        "Melakukan validasi kebutuhan agar fitur yang dibuat sesuai ruang lingkup proyek.",
        "Mendukung koordinasi progres dan pengecekan kesesuaian hasil kerja tim.",
    ]:
        bullet(doc, item)

    doc.add_heading("b. Hasil Kerja Berupa Capture dan Penjelasan", level=3)
    doc.add_picture(str(ERD), width=Inches(6.45))
    caption(doc, "Capture Arico - ERD Sikasir-4SR sebagai hasil analisis relasi data menu, kategori, meja, order, dan detail order.")
    doc.add_paragraph(
        "Capture ERD menunjukkan hubungan satu-ke-banyak antara categories dan menu_items, tables dan orders, orders dan order_items, "
        "serta menu_items dan order_items. ERD ini membantu memastikan alur menu-pesanan sesuai dengan kebutuhan sistem kasir."
    )

    doc.add_heading("c. Hasil Kerja Berupa Coding/Skrip dan Penjelasan", level=3)
    codeblock(
        doc,
        "Relasi model yang mendukung ERD",
        """
class MenuItem extends Model
{
    public function category()
    {
        return $this->belongsTo(Category::class);
    }
}

class Order extends Model
{
    public function order_items()
    {
        return $this->hasMany(OrderItem::class);
    }

    public function table()
    {
        return $this->belongsTo(RestaurantTable::class, 'table_id');
    }
}
""",
        "Potongan kode ini menjadi bukti implementasi relasi yang dirancang pada ERD. Menu terhubung ke kategori, order memiliki banyak item pesanan, dan order terhubung ke nomor meja.",
    )

    doc.add_heading("5.2 RENO SAPUTRA - Front-end dan User Interface", level=2)
    doc.add_heading("a. Daftar/List Pekerjaan Sesuai JobDesc", level=3)
    for item in [
        "Mendesain dan menyesuaikan tampilan antarmuka aplikasi berdasarkan kebutuhan sistem dan referensi Figma.",
        "Mengerjakan tampilan halaman menu/customer, login admin, dashboard admin, barang dan stok, riwayat transaksi, laporan, dan pengaturan.",
        "Menyusun struktur layout, sidebar, topbar, kartu produk, tabel data, metrik dashboard, dan komponen visual.",
        "Menyesuaikan tampilan agar alur kasir dan admin mudah dipahami pengguna.",
        "Menyediakan capture UI di folder proyek-ui sebagai bukti hasil tampilan aplikasi.",
    ]:
        bullet(doc, item)

    doc.add_heading("b. Hasil Kerja Berupa Capture dan Penjelasan", level=3)
    for fname, text in [
        ("Menu Kasir.png", "Capture Reno - Halaman kasir menampilkan daftar produk, kategori, stok, keranjang, nomor meja, dan ringkasan pembayaran."),
        ("Menu Barang & Stok.png", "Capture Reno - Halaman barang dan stok menampilkan tabel inventori, status aktif, stok minimum, dan tombol tambah/edit barang."),
        ("Menu Laporan.png", "Capture Reno - Halaman laporan menampilkan metrik total transaksi, pendapatan, pengeluaran, laba kotor, dan visualisasi laporan."),
    ]:
        img = ROOT / "proyek-ui" / fname
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.35))
            caption(doc, text)
    doc.add_paragraph(
        "Capture tersebut menunjukkan hasil penyesuaian UI agar pengguna dapat melihat informasi penting secara cepat. Halaman kasir fokus pada proses transaksi, "
        "sedangkan halaman admin fokus pada pengelolaan data dan pemantauan laporan."
    )

    doc.add_heading("c. Hasil Kerja Berupa Coding/Skrip dan Penjelasan", level=3)
    codeblock(
        doc,
        "Struktur tampilan kartu produk dan tombol tambah",
        """
<article class="product-card" data-name="{{ strtolower($menu->name) }}" data-category="{{ $menu->category_id }}">
    <div class="product-image">
        <span class="product-badge">{{ $menu->category?->name ?? 'Menu' }}</span>
    </div>
    <button type="button" class="add-button"
        data-id="{{ $menu->id }}"
        data-name="{{ $menu->name }}"
        data-price="{{ $menu->price }}"
        data-stock="{{ $stock }}">+</button>
</article>
""",
        "Potongan Blade ini mewakili pekerjaan UI pada halaman kasir. Data produk disiapkan sebagai atribut HTML agar JavaScript dapat membaca nama, harga, stok, dan id produk saat tombol tambah ditekan.",
    )
    codeblock(
        doc,
        "Interaksi keranjang pada JavaScript",
        """
const subtotal = getCartItems().reduce((total, item) => total + item.price * item.quantity, 0);
const tax = Math.round(taxable * 0.11);
const total = taxable + tax;
inputs.insertAdjacentHTML('beforeend', `<input type="hidden" name="items[${index}][id]" value="${item.id}">`);
""",
        "Script ini mendukung interaksi front-end: menghitung subtotal, PPN, total, serta membuat input hidden supaya item keranjang dapat dikirim ke backend.",
    )

    doc.add_heading("5.3 MUSTOFA ABDURRAHIM - Back-end Laravel, Database, dan Dokumentasi Teknis", level=2)
    doc.add_heading("a. Daftar/List Pekerjaan Sesuai JobDesc", level=3)
    for item in [
        "Mengembangkan Back-end Laravel melalui route, controller, model, dan struktur MVC aplikasi.",
        "Mengelola database MySQL/MariaDB untuk menu, stok, pesanan, user, dan relasi transaksi.",
        "Membuat login admin, CRUD menu, validasi stok, penyimpanan pesanan, serta update status pesanan.",
        "Menggunakan validasi Laravel, DB::transaction, dan lockForUpdate agar proses order lebih aman.",
        "Menyusun dokumentasi progres, catatan teknis, dan laporan portofolio proyek.",
    ]:
        bullet(doc, item)

    doc.add_heading("b. Hasil Kerja Berupa Capture dan Penjelasan", level=3)
    for fname, text in [
        ("Menu Riwayat Transaksi.png", "Capture Mustofa - Riwayat transaksi menampilkan data order, total item, total pembayaran, metode bayar, dan status."),
        ("Menu Pengaturan.png", "Capture Mustofa - Halaman pengaturan menjadi bagian administrasi sistem yang terhubung dengan akses admin."),
    ]:
        img = ROOT / "proyek-ui" / fname
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.35))
            caption(doc, text)
    doc.add_paragraph(
        "Capture ini menunjukkan hasil integrasi data backend ke halaman admin. Data order yang disimpan melalui controller dapat ditampilkan sebagai riwayat transaksi dan dasar laporan."
    )

    doc.add_heading("c. Hasil Kerja Berupa Coding/Skrip dan Penjelasan", level=3)
    codeblock(
        doc,
        "Route utama aplikasi (routes/web.php)",
        """
Route::get('/', [PosController::class, 'index'])->name('pos.index');
Route::post('/orders', [OrderController::class, 'store'])->name('orders.store');
Route::middleware('auth')->group(function () {
    Route::get('/admin', [AdminController::class, 'inventory'])->name('admin.index');
    Route::post('/admin/menu', [AdminController::class, 'storeMenu'])->name('admin.menu.store');
});
""",
        "Route memisahkan halaman publik dan halaman admin. Fitur admin diletakkan di dalam middleware auth sehingga harus login terlebih dahulu.",
    )
    codeblock(
        doc,
        "Validasi dan transaksi order (OrderController)",
        """
$validated = $request->validate([
    'table_id' => ['required', 'exists:tables,id'],
    'items' => ['required', 'array', 'min:1'],
    'items.*.id' => ['required', 'exists:menu_items,id'],
    'items.*.quantity' => ['required', 'integer', 'min:1'],
]);

$order = DB::transaction(function () use ($validated) {
    $menus = MenuItem::whereIn('id', $menuIds)->lockForUpdate()->get()->keyBy('id');
    // hitung subtotal, buat order, buat order_items, kurangi stok
});
""",
        "Backend tetap melakukan validasi walaupun data berasal dari UI. Transaksi database dan lockForUpdate menjaga stok agar tetap konsisten.",
    )
    doc.add_heading("6. Lampiran Pendukung: Pengujian, Temuan, dan Struktur File", level=1)
    doc.add_heading("6.1 Pengujian dan Hasil Pemeriksaan", level=2)
    table(
        doc,
        ["Aktivitas", "Hasil", "Catatan"],
        [
            ["php artisan route:list", "Berhasil", "Terdapat 22 route aktif, termasuk POS, login, order, dan admin."],
            ["php artisan test", "Berhasil", "2 test bawaan Laravel lulus. Test masih minimal sehingga perlu ditambah untuk order dan admin."],
            ["php artisan migrate:status", "Gagal koneksi DB", "MySQL 127.0.0.1:3306 menolak koneksi saat pengecekan. Jalankan MySQL/XAMPP terlebih dahulu."],
            ["Review skema migration", "Perlu perbaikan", "Migration belum selaras dengan model/controller aktif: category_id, is_available, total_price, table_id belum tercermin penuh."],
        ],
        [1.75, 1.3, 3.4],
        8.5,
    )

    doc.add_heading("6.2 Temuan Teknis dan Rekomendasi", level=2)
    for item in [
        "Selaraskan migration dengan model/controller aktif. Tambahkan tabel categories dan tables, ubah menu_items agar memiliki category_id, description, is_available, dan pastikan orders memakai table_id serta total_price.",
        "Tambahkan kolom subtotal, tax, discount_type, discount_value, discount_amount, total_price, dan payment_method jika diskon dan metode bayar ingin tersimpan secara akurat.",
        "Aktifkan kembali route CustomerMenuController jika fitur pelanggan via QR meja ingin digunakan. Saat ini route /menu dan /menu/meja/{number} redirect ke POS.",
        "Tambahkan feature test untuk login admin, CRUD menu, pembuatan order, validasi stok habis, dan perubahan status pesanan.",
        "Perbarui README agar tidak lagi berisi template Laravel default, tetapi berisi deskripsi Sikasir-4SR, instalasi, akun demo, database, dan screenshot.",
    ]:
        num(doc, item)

    doc.add_heading("6.3 Kesimpulan", level=2)
    doc.add_paragraph(
        "Portofolio ini menunjukkan bahwa proyek Sikasir-4SR telah memiliki fondasi aplikasi kasir berbasis Laravel yang mencakup "
        "antarmuka kasir, pengelolaan barang dan stok, autentikasi admin, pemrosesan order, riwayat transaksi, serta laporan penjualan. "
        "Dari sisi pembelajaran, proyek telah mencakup aktivitas Front-end, API/route, dan Back-end sebagaimana diminta pada ketentuan portofolio. "
        "Tahap berikutnya adalah menyelaraskan migration database, memperluas test, dan melengkapi data identitas tim sebelum laporan dicetak atau dikumpulkan."
    )

    doc.add_heading("6.4 Struktur File yang Relevan", level=2)
    table(
        doc,
        ["Area", "File/Folder", "Keterangan"],
        [
            ["Route", "routes/web.php", "Definisi URL aplikasi dan middleware."],
            ["Controller", "app/Http/Controllers", "Logika POS, order, admin, auth, customer menu."],
            ["Model", "app/Models", "Representasi data user, menu, kategori, meja, order, order item."],
            ["View", "resources/views", "Blade template untuk UI kasir, admin, login, home."],
            ["Asset", "public/assets/css/app.css", "Styling utama aplikasi."],
            ["Asset", "public/assets/js/pos.js", "Interaksi keranjang, filter, kalkulasi total, modal barang."],
            ["Database", "database/migrations", "Definisi struktur tabel."],
            ["Seeder", "database/seeders/MenuItemSeeder.php", "Data awal menu."],
            ["Dokumentasi", "docs/struktur-laravel.md", "Catatan setup dan struktur Laravel."],
        ],
        [1.0, 2.35, 3.1],
        8.5,
    )

    doc.save(DOCX)
    print(DOCX)
    print(ERD)


if __name__ == "__main__":
    main()
