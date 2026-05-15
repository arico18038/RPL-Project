from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Laporan_Penggunaan_Tools_Cash_Dig.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    set_paragraph_spacing(paragraph, after=4)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)
    set_paragraph_spacing(paragraph, after=4)
    return paragraph


def add_tools_table(document, rows):
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    widths = [1800, 2200, 3100, 2260]
    headers = ["No.", "Tools", "Kegunaan", "Alasan Penggunaan"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    set_repeat_table_header(table.rows[0])

    for number, tool, use, reason in rows:
        cells = table.add_row().cells
        values = [str(number), tool, use, reason]
        for index, value in enumerate(values):
            set_cell_width(cells[index], widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            set_paragraph_spacing(paragraph, after=0, line=1.1)

    return table


def add_signature_block(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, before=18, after=4)
    paragraph.add_run("Penyusun,\n\n\n")
    run = paragraph.add_run("(Nama Mahasiswa yang Submit)")
    run.bold = True


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=40, after=12)
    run = title.add_run("LAPORAN PENGGUNAAN TOOLS\nPENGEMBANGAN APLIKASI CASH-DIG")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=34)
    run = subtitle.add_run("Disusun untuk memenuhi tugas mata kuliah pengembangan perangkat lunak")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    cover_items = [
        ("Nama Kelompok", "Cash-Dig / 4SR"),
        ("Anggota 1", "Nama Anggota 1"),
        ("Anggota 2", "Nama Anggota 2"),
        ("Nama yang Submit", "NAMA MAHASISWA YANG SUBMIT"),
        ("Kelas", "Isi kelas"),
        ("Program Studi", "Isi program studi"),
        ("Tahun", "2026"),
    ]

    cover_table = document.add_table(rows=0, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.autofit = False
    set_table_borders(cover_table)
    for label, value in cover_items:
        row = cover_table.add_row()
        set_cell_width(row.cells[0], 2600)
        set_cell_width(row.cells[1], 5000)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "F2F4F7")
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        value_run = row.cells[1].paragraphs[0].add_run(value)
        if label == "Nama yang Submit":
            value_run.bold = True
            value_run.underline = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    document.add_page_break()

    add_heading(document, "1. Pendahuluan", 1)
    paragraph = document.add_paragraph(
        "Cash-Dig adalah aplikasi kasir dan pemesanan menu berbasis web yang dirancang untuk membantu proses transaksi pada rumah makan. "
        "Aplikasi ini memiliki dua sisi utama, yaitu dashboard admin/kasir untuk mengelola barang, stok, pesanan, riwayat, laporan, dan pengaturan, "
        "serta halaman pelanggan yang dapat diakses melalui barcode atau QR code pada meja."
    )
    set_paragraph_spacing(paragraph)

    add_heading(document, "2. Kesepakatan Tim", 1)
    paragraph = document.add_paragraph(
        "Berdasarkan kesepakatan tim, aplikasi dibangun sebagai aplikasi web menggunakan arsitektur client-server. "
        "Tim memilih Laravel sebagai framework backend karena mendukung routing, controller, validasi, model, dan integrasi database secara terstruktur. "
        "Database yang digunakan adalah MySQL/MariaDB melalui XAMPP karena mudah digunakan pada lingkungan pengembangan lokal."
    )
    set_paragraph_spacing(paragraph)

    add_bullet(document, "Basis aplikasi: web application.")
    add_bullet(document, "Framework utama: Laravel.")
    add_bullet(document, "Database: MySQL/MariaDB.")
    add_bullet(document, "Tools kolaborasi: Git dan GitHub.")
    add_bullet(document, "Editor pengembangan: Visual Studio Code.")

    add_heading(document, "3. Tools yang Digunakan", 1)
    tools = [
        (1, "Laravel", "Framework backend PHP untuk membuat route, controller, model, validasi form, dan tampilan Blade.", "Mempercepat pengembangan aplikasi web dan membuat struktur proyek lebih rapi."),
        (2, "PHP", "Bahasa pemrograman utama untuk menjalankan Laravel.", "Laravel berjalan di atas PHP sehingga diperlukan untuk proses backend."),
        (3, "MySQL/MariaDB", "Database untuk menyimpan data kategori, menu, meja, pesanan, dan detail pesanan.", "Cocok untuk data transaksi yang terstruktur dan mudah dikelola melalui XAMPP."),
        (4, "XAMPP", "Paket lokal server yang menyediakan Apache, PHP, dan MySQL/MariaDB.", "Mempermudah pengembangan dan pengujian aplikasi di komputer lokal."),
        (5, "Blade Template", "Template engine Laravel untuk membuat halaman admin, kasir, dan pelanggan.", "Membantu memisahkan tampilan dari logika aplikasi."),
        (6, "HTML, CSS, JavaScript", "Membangun antarmuka pengguna, styling, modal, cart, serta interaksi halaman.", "Diperlukan agar aplikasi mudah digunakan dan tampil sesuai desain."),
        (7, "Git", "Version control untuk mencatat perubahan kode.", "Memudahkan tim melihat riwayat perubahan dan menghindari kehilangan pekerjaan."),
        (8, "GitHub", "Repository online untuk menyimpan dan membagikan source code.", "Mendukung kolaborasi tim melalui pull, push, dan backup kode."),
        (9, "Visual Studio Code", "Editor kode untuk mengubah file Laravel, Blade, CSS, dan JavaScript.", "Ringan, banyak ekstensi, dan nyaman untuk pengembangan web."),
        (10, "Figma", "Tools desain UI sebelum diimplementasikan ke website.", "Membantu tim menyepakati tampilan aplikasi sebelum dibuat di kode."),
    ]
    add_tools_table(document, tools)

    add_heading(document, "4. Pembagian Fungsi Aplikasi", 1)
    add_heading(document, "4.1 Dashboard Admin/Kasir", 2)
    paragraph = document.add_paragraph(
        "Dashboard admin digunakan untuk mengelola barang dan stok, melihat pesanan masuk, memproses pesanan, menyelesaikan pesanan, melihat riwayat transaksi, serta membuka laporan dan pengaturan."
    )
    set_paragraph_spacing(paragraph)

    add_heading(document, "4.2 Halaman Pelanggan", 2)
    paragraph = document.add_paragraph(
        "Pelanggan duduk pada meja yang memiliki nomor, kemudian melakukan scan barcode/QR code. Barcode tersebut mengarahkan pelanggan ke halaman menu berbasis web. "
        "Pelanggan dapat memilih makanan atau minuman, memasukkan pesanan ke keranjang, lalu mengirim pesanan ke sistem."
    )
    set_paragraph_spacing(paragraph)

    add_heading(document, "5. Alur Pengembangan", 1)
    workflow = [
        "Tim membuat rancangan tampilan menggunakan Figma.",
        "Desain diimplementasikan menjadi halaman Laravel menggunakan Blade, CSS, dan JavaScript.",
        "Database MySQL disiapkan untuk menyimpan kategori, menu, meja, pesanan, dan item pesanan.",
        "Fitur admin dibuat untuk menambah, mengedit, menghapus menu, serta memproses pesanan.",
        "Fitur pelanggan dibuat agar pesanan dapat dikirim dari halaman menu berbasis barcode meja.",
        "Source code disimpan dan diperbarui menggunakan Git serta GitHub.",
    ]
    for item in workflow:
        add_number(document, item)

    add_heading(document, "6. Kelebihan Penggunaan Tools", 1)
    add_bullet(document, "Laravel membuat struktur kode lebih terorganisir melalui konsep route, controller, model, dan view.")
    add_bullet(document, "MySQL/MariaDB memudahkan penyimpanan data transaksi dan relasi antar tabel.")
    add_bullet(document, "GitHub membantu tim menyimpan kode secara online dan mengambil perubahan terbaru dari anggota lain.")
    add_bullet(document, "Figma membantu menyamakan persepsi tampilan sebelum proses implementasi.")
    add_bullet(document, "VS Code mempercepat proses pengeditan kode karena mendukung banyak ekstensi web development.")

    add_heading(document, "7. Kesimpulan", 1)
    paragraph = document.add_paragraph(
        "Pemilihan Laravel, MySQL/MariaDB, XAMPP, Git, GitHub, Visual Studio Code, dan Figma dinilai sesuai untuk pengembangan aplikasi Cash-Dig. "
        "Kombinasi tools tersebut mendukung kebutuhan pengembangan aplikasi kasir dan pemesanan menu berbasis web, mulai dari perancangan tampilan, pembuatan fitur, pengelolaan database, sampai kolaborasi kode antar anggota tim."
    )
    set_paragraph_spacing(paragraph)

    add_signature_block(document)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
