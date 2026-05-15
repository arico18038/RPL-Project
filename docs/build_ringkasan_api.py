from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Ringkasan_API_dan_RESTful_API.docx"


def set_spacing(paragraph, before=0, after=8, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_run_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(paragraph, after=8)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    set_run_font(run)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=10, after=14)
    run = title.add_run("Ringkasan Materi API dan RESTful API")
    set_run_font(run, size=16, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(meta, after=18)
    run = meta.add_run("Sumber: Hostinger Tutorial")
    set_run_font(run, size=10)

    heading = document.add_heading("Ringkasan", level=1)
    for run in heading.runs:
        set_run_font(run, size=14, bold=True)

    add_body(
        document,
        "API atau Application Programming Interface adalah penghubung yang membuat satu aplikasi bisa berkomunikasi dengan aplikasi lain. "
        "API berperan seperti perantara: aplikasi meminta data atau fitur tertentu, lalu API meneruskan permintaan tersebut ke server dan mengembalikan hasilnya ke aplikasi. "
        "Dengan API, developer tidak perlu membuat semua fitur sendiri dari awal. Contohnya, aplikasi transportasi bisa memakai API Google Maps untuk menampilkan peta."
    )

    add_body(
        document,
        "API bermanfaat untuk mempercepat pengembangan aplikasi, menghemat biaya, memudahkan integrasi antarplatform, dan mengurangi beban server. "
        "API juga memiliki beberapa jenis, seperti Public API yang bisa digunakan umum, Private API untuk kebutuhan internal, Partner API untuk pihak yang bekerja sama, "
        "dan Composite API yang menggabungkan data dari beberapa sumber."
    )

    add_body(
        document,
        "RESTful API adalah API yang mengikuti prinsip arsitektur REST atau Representational State Transfer. RESTful API biasanya menggunakan protokol HTTP "
        "dan memakai metode seperti GET untuk mengambil data, POST untuk menambah data, PUT untuk memperbarui data, dan DELETE untuk menghapus data. "
        "Data yang dikirim biasanya berbentuk JSON atau XML."
    )

    add_body(
        document,
        "Ciri utama RESTful API adalah client dan server terpisah, komunikasi bersifat stateless, mendukung cache, memiliki interface yang seragam, "
        "dan bisa menggunakan sistem berlapis. Stateless berarti setiap permintaan harus membawa informasi yang lengkap, sehingga server tidak perlu mengingat permintaan sebelumnya."
    )

    conclusion = document.add_heading("Kesimpulan", level=1)
    for run in conclusion.runs:
        set_run_font(run, size=14, bold=True)

    add_body(
        document,
        "Kesimpulannya, API adalah jembatan komunikasi antar aplikasi, sedangkan RESTful API adalah salah satu cara atau standar dalam membuat API agar pertukaran data lebih rapi, fleksibel, dan mudah digunakan. "
        "API dan RESTful API sangat penting dalam pengembangan aplikasi modern karena membantu sistem saling terhubung dengan lebih efisien."
    )

    sources = document.add_heading("Sumber", level=1)
    for run in sources.runs:
        set_run_font(run, size=14, bold=True)

    add_bullet(document, "https://www.hostinger.com/id/tutorial/api-adalah")
    add_bullet(document, "https://www.hostinger.com/id/tutorial/apa-itu-restful-api")

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
